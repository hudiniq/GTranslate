from googletrans import Translator
import docx
from os import path
import urllib.request
import pypandoc


class Translate():
    def __init__(self):
        doc_final = docx.Document()
        translator = Translator()
        result = []
        self.srcs = ["CN", "BR", "IN", "AUTO"]

        print("Vnesite ime wordove datoteke, oz. URL:")
        while True:
            try:
                self.usersource = self.source_input(input())
            except FileNotFoundError:
                print("\nDatoteka ne obstaja ali je v napačnem formatu.\nProsim za ponovni vnos:")
                continue
            break

        print("Prevod iz katerega jezika?\nNa voljo so CN, BR, IN in AUTO:")
        while True:
            try:
                user_src = self.src_input(input())
            except ValueError:
                print("\nNapačen izvorni jezik:")
                continue
            break
        
        if user_src in self.srcs[:-1]:
            result.append(translator.translate(self.usersource, src=user_src, dest='en'))
        elif user_src == "AUTO":
            result.append(translator.translate(self.usersource, dest='en'))

        doc_final.add_paragraph("Translated from " + result[0].src + " to " + result[0].dest + ".")
        for par in result:
            doc_final.add_paragraph(par.text)

        doc_final.save("translated.docx")

    def source_input(self, source):
        try:
            html = urllib.request.urlopen(source).read()
        except ValueError:
            if ".docx" not in source:
                source = source + ".docx"

            if not path.exists(source):
                raise FileNotFoundError

            return self.text_from_doc(source)

        return self.doc_from_html(html)

    def doc_from_html(self, html):
        while True:
            try:
                pypandoc.convert_text(source=html, format='html', to='docx', outputfile='html_pre_translate.docx', extra_args=['-RTS'])
            except OSError:
                pypandoc.download_pandoc()
                continue
            break
        return self.text_from_doc("html_pre_translate.docx")

    def text_from_doc(self, filename):
        doc = docx.Document(filename)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)

    def src_input(self, src):
        src = str.upper(src)

        if src not in self.srcs:
            raise ValueError

        if src == "CN":
            src = "ZN-CN"

        return src


if __name__ == "__main__":
    app = Translate()